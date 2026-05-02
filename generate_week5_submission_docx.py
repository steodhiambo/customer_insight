from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def set_font(run, size, bold=False, color=BLACK, italic=False):
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = color

def heading(text, level=1, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 14 if level == 1 else 12, bold=True, color=color)
    return p

def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, 11)
    return p

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.width = Inches(col_widths[i])
        shade_cell(cell, "D9D9D9")
        run = cell.paragraphs[0].add_run(h)
        set_font(run, 10, bold=True, color=BLACK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.width = Inches(col_widths[c_idx])
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, 10, color=BLACK)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ═════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Travler — Week 5: Booking Journey & Conversion Analysis")
set_font(run, 18, bold=True)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Submission by Stephen")
set_font(run2, 12, italic=True)
p2.paragraph_format.space_after = Pt(20)

# ═════════════════════════════════════════════════════════════════════════════
# TASK 1
# ═════════════════════════════════════════════════════════════════════════════
heading("Task 1: Funnel Analysis (Drop-off Points)")
body(
    "The funnel data reveals significant friction at key transitions. Here are the 3 key drop-off points identified:"
)

drop_offs = [
    ["Payment Page Entry", "37% drop-off from previous stage (4,800 to 3,000 users).", "Suggests a critical trust barrier or technical failure just as users are asked to commit financial details."],
    ["Route Selection", "27% drop-off from search (8,500 to 6,200 users).", "Users are searching but not finding options that meet their price or value expectations, leading to external comparison."],
    ["Passenger Details Entry", "22% drop-off from selection (6,200 to 4,800 users).", "Suggests 'form fatigue'. Users are interested in the route but find the data entry process too cumbersome to start."]
]

add_table(["Drop-off Location", "Data Observation", "User Behaviour Insight"], drop_offs, [1.5, 2.0, 3.0])

# ═════════════════════════════════════════════════════════════════════════════
# TASK 2
# ═════════════════════════════════════════════════════════════════════════════
heading("Task 2: Friction Analysis")
body(
    "The following friction areas are suppressing conversion by undermining user trust and ease of use:"
)

friction_areas = [
    ["Payment Security (Trust)", "Users are unsure if the platform is secure (F001). This creates high anxiety at the final hurdle, leading to the 37% drop-off.", "Payment Page"],
    ["External Price Comparison (Value)", "Users leave to compare prices (F002). This indicates that Travler's value proposition is not immediately clear or convincing.", "Route Selection / Payment"],
    ["Complex Forms (Usability)", "Too many required fields (F003) frustrate users, particularly on mobile where bounce rates are high (up to 55%).", "Passenger Details Entry"],
    ["System Reliability (Trust)", "Reported payment failures (F004) damage brand reputation and cause immediate abandonment of the booking journey.", "Payment Page"],
    ["Inventory Transparency (Usability)", "Unclear seat updates (F006) cause hesitation. Users need certainty about availability before committing to enter details.", "Route Selection"]
]

add_table(["Friction Area", "Impact on User Decision", "Journey Stage"], friction_areas, [1.8, 3.2, 1.5])

# ═════════════════════════════════════════════════════════════════════════════
# TASK 3
# ═════════════════════════════════════════════════════════════════════════════
heading("Task 3: Improvement Recommendations")
body(
    "To address these issues and improve conversion, the following practical changes are recommended:"
)

recommendations = [
    ["Visible Trust Signals", "Payment Security (F001)", "Adding security badges and partner logos (e.g., M-Pesa) provides visual reassurance, directly tackling the 37% drop-off."],
    ["Streamlined Detail Entry", "Form Fatigue (F003)", "Implementing social login or 'guest' checkout with minimal fields reduces friction, especially for mobile-first segments."],
    ["Transparent 'All-In' Pricing", "Price Comparison (F002)", "Showing final prices early and adding a 'Best Price Guarantee' badge reduces the urge for users to check external sites."],
    ["Real-Time Availability Alerts", "Seat Uncertainty (F006)", "Using urgency cues (e.g., 'Only 2 seats left') clarifies inventory and encourages users to proceed to the next stage."]
]

add_table(["Recommendation", "Issue Addressed", "Why it Improves Conversion"], recommendations, [1.5, 1.5, 3.5])

# ═════════════════════════════════════════════════════════════════════════════
# TASK 4
# ═════════════════════════════════════════════════════════════════════════════
heading("Task 4: Prioritisation of Actions")
body(
    "The top two improvements with the highest potential for impact are:"
)

priorities = [
    ["1. Security Trust Signals", "This addresses the largest drop-off (37%) at the most critical stage. Trust is non-negotiable for online payments; solving this recovers the most lost revenue."],
    ["2. Form Simplification", "This directly benefits the high-volume, high-bounce mobile segments. Reducing entry barriers is the fastest way to improve flow for the majority of Travler's traffic."]
]

add_table(["Priority Improvement", "Justification for Priority"], priorities, [2.0, 4.5])

doc.save("Stephen_Travler_Week5.docx")
print("Saved -> Stephen_Travler_Week5.docx")
