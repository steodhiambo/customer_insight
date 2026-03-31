from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK  = RGBColor(0x00, 0x00, 0x00)
BLUE  = RGBColor(0x00, 0x00, 0x00)
RED   = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0xF4, 0xF6, 0xF8)

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Style helpers ────────────────────────────────────────────────────────────
def set_font(run, size, bold=False, color=BLACK, italic=False):
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    run.font.color.rgb = color

def heading(text, level=1, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    size = 16 if level == 1 else 12
    set_font(run, size, bold=True, color=color)
    return p

def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, 10)
    return p

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths, header_hex="000000"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.width = Inches(col_widths[i])
        shade_cell(cell, "FFFFFF")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        set_font(run, 9, bold=True, color=BLACK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # data rows
    for r_idx, row in enumerate(rows):
        bg = "F4F6F8" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.width = Inches(col_widths[c_idx])
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            # colour conversion column values
            if headers[c_idx] in ("Conversion Rate", "Conversion") and "%" in str(val):
                set_font(run, 9, bold=True, color=BLACK)
            else:
                set_font(run, 9, color=BLACK)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Travler — Week 1 Customer Insight Analysis")
set_font(r, 22, bold=True, color=BLACK)

p2 = doc.add_paragraph()
r2 = p2.add_run("Marketing Team Submission  |  Behavioural Analysis & Segmentation")
set_font(r2, 11, color=BLACK, italic=True)
p2.paragraph_format.space_after = Pt(10)

# ════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════
heading("1. Executive Summary")
body(
    "Travler does not have a traffic problem. The platform is attracting users who actively "
    "search routes and compare prices. The problem is conversion — only 45% of users "
    "(9 of 20) completed a booking. The gap between traffic and bookings is explained by "
    "three compounding factors: mobile checkout friction, price uncertainty, and a mismatch "
    "between platform design and the behaviour of the platform's largest user groups."
)

add_table(
    ["Metric", "Value", "Implication"],
    [
        ["Overall Conversion Rate",        "45%  (9 / 20 users)", "More than half of all users leave without booking"],
        ["Desktop Conversion",             "86%",                 "Desktop experience is working well"],
        ["Mobile Conversion",              "23%",                 "Mobile checkout is the single biggest drop-off point"],
        ["Morning / Afternoon Conversion", "83 – 100%",           "Focused users in low-distraction windows convert reliably"],
        ["Evening / Night Conversion",     "0%",                  "Distracted browsing sessions never complete"],
        ["Avg. Searches — Converters",     "1.7",                 "Low frequency = clear intent"],
        ["Avg. Searches — Non-converters", "4.5",                 "High frequency = price uncertainty and hesitation"],
    ],
    [2.2, 1.3, 2.8]
)

# ════════════════════════════════════════════════════════════════════════════
# 2. BEHAVIOURAL PATTERNS
# ════════════════════════════════════════════════════════════════════════════
heading("2. Behavioural Patterns")
body(
    "The dataset was analysed across five dimensions — travel purpose, age group, device, "
    "time of search, and search frequency — to identify which factors consistently predict "
    "whether a user books or abandons."
)

heading("2.1  Conversion by Travel Purpose", level=2, color=BLACK)
add_table(
    ["Travel Purpose", "Conversion Rate", "Interpretation"],
    [
        ["Family",   "100%", "Committed and purposeful — searches once and books"],
        ["Business", "83%",  "Clear intent, time-sensitive, low tolerance for friction"],
        ["Leisure",  "20%",  "Interested but not urgent — comparing options"],
        ["Student",  "0%",   "Highest search volume on the platform, zero bookings"],
    ],
    [1.5, 1.3, 3.5]
)

heading("2.2  Device Gap — The Biggest Single Finding", level=2, color=BLACK)
body(
    "Desktop users convert at 86%. Mobile users convert at 23%. This is not a content or "
    "pricing problem — it is a checkout experience problem. The payment form is harder to "
    "complete on a small screen, trust signals disappear on mobile layouts, and any "
    "interruption during a mobile session results in permanent abandonment. "
    "76% of non-converters are on mobile."
)

heading("2.3  Time of Search", level=2, color=BLACK)
add_table(
    ["Time of Search", "Conversion Rate", "Context"],
    [
        ["Morning",   "83%",  "Focused and intentional — likely at a desk"],
        ["Afternoon", "100%", "Peak conversion window — no distractions"],
        ["Evening",   "0%",   "Browsing from mobile in a distracted environment"],
        ["Night",     "0%",   "Late-night comparison shopping — never completes"],
    ],
    [1.5, 1.3, 3.5]
)

heading("2.4  Search Frequency as a Hesitation Signal", level=2, color=BLACK)
body(
    "Users who booked averaged 1.7 searches. Users who did not book averaged 4.5 searches. "
    "High search frequency does not signal growing intent — it signals that the user cannot "
    "find a reason to stop searching. The most common cause is price uncertainty: the user "
    "does not trust that the price they see is stable or the best available, so they keep "
    "checking. This is the same concern Brian O.'s support team is fielding on customer calls."
)

# ════════════════════════════════════════════════════════════════════════════
# 3. CUSTOMER SEGMENTS
# ════════════════════════════════════════════════════════════════════════════
heading("3. Customer Segments")
body(
    "Four distinct behavioural segments were identified by combining travel purpose, device, "
    "time of search, and search frequency. Each segment has a consistent profile that repeats "
    "across multiple users in the dataset."
)

add_table(
    ["Segment", "Age", "Device", "Time", "Searches", "Conversion"],
    [
        ["Decisive Business Traveller", "25–44", "Desktop", "Morning",   "1–3", "83%"],
        ["Family Planner",              "45+",   "Mixed",   "Afternoon", "1",   "100%"],
        ["Casual Leisure Browser",      "25–44", "Mobile",  "Evening",   "2–3", "20%"],
        ["Price-Sensitive Student",     "18–24", "Mobile",  "Night",     "5–7", "0%"],
    ],
    [2.0, 0.7, 0.9, 1.0, 1.0, 1.0]
)

segments = [
    ("Decisive Business Traveller — 83% conversion",
     "Knows the route, has a deadline, and books with minimal friction. Uses desktop in the "
     "morning and searches 1–3 times before committing. Motivation: efficiency and reliability. "
     "Barrier: any friction in the checkout flow."),
    ("Family Planner — 100% conversion",
     "The most reliable converter on the platform. Searches once in the afternoon and commits "
     "immediately. Motivation: certainty and trust. Barrier: confusion or complexity in the "
     "booking process."),
    ("Casual Leisure Browser — 20% conversion",
     "Interested but not urgent. Browses on mobile in the evening while multitasking. Compares "
     "2–3 options but rarely commits. Motivation: finding the best deal. Barrier: mobile "
     "checkout friction and no urgency trigger."),
    ("Price-Sensitive Student Searcher — 0% conversion",
     "The most active searcher on the platform — 5 to 7 searches per session — but never books. "
     "Uses mobile at night. Motivation: travel intent is real. Barrier: price is the hard block. "
     "No discount, no pay-later option means no booking."),
]
for name, desc in segments:
    heading(name, level=2, color=BLACK)
    body(desc)

# ════════════════════════════════════════════════════════════════════════════
# 4. MOTIVATIONS AND BARRIERS
# ════════════════════════════════════════════════════════════════════════════
heading("4. Motivations and Barriers")
add_table(
    ["Factor", "Motivations (drives booking)", "Barriers (blocks booking)"],
    [
        ["Device",            "Desktop: full keyboard, autofill, trust signals visible",
                              "Mobile: small screen, autofill fails, trust signals hidden"],
        ["Time of Search",    "Morning/Afternoon: focused, intentional sessions",
                              "Evening/Night: distracted environment, sessions abandoned"],
        ["Search Frequency",  "1–2 searches: clear intent, low uncertainty",
                              "5–7 searches: price uncertainty, no confidence mechanism"],
        ["Travel Purpose",    "Business/Family: clear need, deadline, or commitment",
                              "Student/Leisure: price-sensitive, no urgency, no trust anchor"],
        ["Pricing",           "Stable, transparent pricing builds confidence to book",
                              "Inconsistent pricing triggers repeat searching without booking"],
    ],
    [1.2, 2.7, 2.7]
)

# ════════════════════════════════════════════════════════════════════════════
# 5. RECOMMENDATIONS
# ════════════════════════════════════════════════════════════════════════════
heading("5. Recommendations")

recs = [
    ("1. Fix Mobile Checkout",
     "76% of non-converters are on mobile. Simplify the payment form, increase tap target sizes, "
     "enable autofill for card details, and ensure trust signals remain visible on small screens. "
     "Add a progress indicator so users know how close they are to finishing."),
    ("2. Retarget Evening and Night Abandoners",
     "Conversion at these times is 0%. Trigger an automated follow-up — push notification or "
     "email — within 1 hour of an abandoned session. The user had intent; they just did not "
     "complete in that window."),
    ("3. Unlock the Student Segment",
     "Students are the most active searchers on the platform with 0% conversion. A student "
     "discount or pay-later option directly removes the price barrier and represents the largest "
     "untapped conversion opportunity on the platform."),
    ("4. Protect Business and Family Users",
     "These segments already convert at 83–100%. Keep the desktop experience frictionless and "
     "introduce loyalty perks — saved routes, priority booking — to increase repeat visits and "
     "lifetime value."),
    ("5. Add Pricing Transparency",
     "Repeated searches on the same route signal that users do not trust the price is stable. "
     "A Best Price Guarantee badge or fare-alert feature converts browsing into booking by "
     "removing the reason to keep searching. This also directly reduces the volume of pricing "
     "inquiries reaching Brian O.'s support team."),
]
for title, desc in recs:
    heading(title, level=2, color=BLACK)
    body(desc)

# ════════════════════════════════════════════════════════════════════════════
# 6. CONCLUSION
# ════════════════════════════════════════════════════════════════════════════
heading("6. Conclusion")
body(
    "Travler's growth challenge is not a traffic problem — it is a conversion problem "
    "concentrated in specific, identifiable segments and touchpoints. The data shows precisely "
    "where losses are happening: mobile checkout, evening and night sessions, the student "
    "segment, and price uncertainty at the point of decision. The five recommendations above "
    "are each tied directly to a finding in the data. Addressing them in order of impact — "
    "mobile first, then pricing transparency, then segment-specific interventions — gives the "
    "marketing and operations teams a clear, evidence-based roadmap for improving conversion "
    "without increasing traffic acquisition spend."
)

# ── Footer note ──────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("Travler  |  Week 1 Customer Insight  |  Dataset: 20 users  |  Analysis: search & booking behaviour")
set_font(r, 8, color=BLACK, italic=True)

doc.save("Travler_Week1_Customer_Insight_Submission.docx")
print("Saved → Travler_Week1_Customer_Insight_Submission.docx")
