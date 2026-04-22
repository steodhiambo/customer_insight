from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

def set_font(run, size, bold=False, color=BLACK, italic=False):
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = color

def heading(text, level=1, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 12, bold=True, color=color)
    return p

def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, 10)
    return p

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.width = Inches(col_widths[i])
        shade_cell(cell, "FFFFFF")
        run = cell.paragraphs[0].add_run(h)
        set_font(run, 9, bold=True, color=BLACK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows):
        bg = "F4F6F8" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.width = Inches(col_widths[c_idx])
            shade_cell(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, 9, color=BLACK)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ═════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(p.add_run("Travler — Week 4 Paid Advertising Strategy Under Budget Constraints"), 22, bold=True)

p2 = doc.add_paragraph()
set_font(p2.add_run(
    "Marketing Team Submission  |  Audience Prioritisation, Channel Strategy, Messaging & Budget Allocation"),
    11, italic=True)
p2.paragraph_format.space_after = Pt(10)

# ═════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
heading("1. Executive Summary")
body(
    "Travler enters Week 4 under increased financial scrutiny following the reputational "
    "incident in Week 3 and a leadership decision to reduce marketing spend while reviewing "
    "effectiveness. The Finance team, led by Kevin R., is demanding clear justification for "
    "every marketing shilling spent. At the same time, Amina K. has stressed that maintaining "
    "visibility in a competitive market is non-negotiable. This submission addresses that "
    "tension directly by building a paid advertising strategy that is both defensible to "
    "Finance and effective for Marketing. The strategy is grounded in three weeks of "
    "accumulated data across customer behaviour, content performance, and campaign results. "
    "It identifies the three audience segments most likely to deliver high conversion value, "
    "proposes a three-channel approach with clearly defined roles, develops three messaging "
    "angles tailored to segment behaviour, and recommends a budget allocation that concentrates "
    "spend where the data shows the highest return."
)

heading("Internal Context", level=2)
body(
    "Kevin R. requires that every channel in the plan can demonstrate a measurable link "
    "between spend and bookings. This submission meets that requirement by anchoring the "
    "primary budget allocation to Google Ads, which delivered the highest ROI score (5/5) "
    "and the lowest cost per booking (KES 350–455) in the dataset. Amina K.'s concern about "
    "visibility is addressed by maintaining Instagram as a retargeting and mid-funnel channel "
    "and retaining X at minimum spend for reputation management — a direct response to the "
    "trust and sentiment issues that remain unresolved from Week 3. TikTok is paused under "
    "this strategy, a decision supported by its Week 3 CPA of KES 1,000–1,125 and ROI "
    "score of 2/5, which cannot be justified under current budget constraints."
)

add_table(
    ["Metric", "Value", "Implication"],
    [
        ["Audience Segments Prioritised",   "3 of 5",                      "Business Travellers, Families, Young Professionals"],
        ["Primary Channel",                 "Google Ads",                  "ROI score 5/5, lowest cost per booking, lowest risk"],
        ["Secondary Channel",               "Instagram",                   "Retargeting and mid-funnel, ROI score 4/5"],
        ["Reputation Channel",              "X — minimum spend",           "Trust recovery post Week 3 incident, ROI score 3/5"],
        ["Channel Paused",                  "TikTok",                      "CPA KES 1,000–1,125, ROI score 2/5 — indefensible under constraints"],
        ["Recommended Budget Split",        "55% / 30% / 15%",             "Google Ads / Instagram / X"],
        ["Segments Excluded",               "Students",                    "1.5% conversion, KES 900 cost per booking — lowest ROI in dataset"],
        ["Top Messaging Priority",          "Trust + Efficiency",          "Addresses both Week 3 sentiment damage and high-value segment needs"],
    ],
    [2.4, 1.8, 2.6]
)

# ═════════════════════════════════════════════════════════════════════════════
# 2. TASK 1 — AUDIENCE SEGMENTS
# ═════════════════════════════════════════════════════════════════════════════
heading("2. Task 1 — Audience Segments Most Likely to Deliver High Conversion Value")
body(
    "The Week 4 Audience Performance dataset, combined with Week 1 behavioural data and "
    "Week 3 campaign results, identifies three segments as the highest-priority targets "
    "under budget constraints. Prioritisation is based on conversion rate, cost per booking, "
    "and volume potential considered together — not any single metric in isolation."
)

add_table(
    ["Segment", "Conv Rate", "Cost per Booking", "Volume", "Platform", "Priority"],
    [
        ["Business Travellers", "14.0%", "KES 350", "Medium", "Google Ads", "1 — Anchor segment"],
        ["Families",            "10.5%", "KES 400", "Low",    "Google Ads", "2 — High ROI, low volume"],
        ["Young Professionals", "6.0%",  "KES 500", "High",   "Instagram",  "3 — High volume, manageable CPA"],
        ["Leisure Travellers",  "4.5%",  "KES 600", "Medium", "Instagram",  "Secondary — retained at low spend"],
        ["Students",            "1.5%",  "KES 900", "High",   "TikTok",     "Excluded — indefensible CPA"],
    ],
    [1.8, 0.9, 1.3, 0.8, 1.1, 1.9]
)

segments = [
    ("Segment 1 — Business Travellers (Priority 1)",
     "Business Travellers are the anchor segment for this strategy. They convert at 14% — "
     "the highest rate in the dataset — at a cost of KES 350 per booking, the lowest in the "
     "dataset. In Week 3, C006 (Google Ads, Business Travellers) delivered 110 bookings from "
     "10,000 impressions at KES 50,000, and C004 (X, Business Travellers) delivered 70 "
     "bookings at KES 30,000 — a CPA of KES 429, the best in the entire Week 3 dataset. "
     "This segment is time-sensitive, decision-ready, and responds to search-based targeting "
     "because they are in an active planning state when they encounter Travler ads. Under "
     "budget constraints, this is the segment that delivers the most bookings per shilling "
     "spent and must receive protected budget allocation."),

    ("Segment 2 — Families (Priority 2)",
     "Families convert at 10.5% at KES 400 per booking — the second-best combination of "
     "conversion rate and cost efficiency in the dataset. In Week 1, the Family Planner "
     "segment achieved a 100% conversion rate, searching once and committing immediately. "
     "Volume is low, which limits total booking numbers, but the cost efficiency makes every "
     "booking from this segment highly profitable. Families respond to trust signals, safety "
     "reassurance, and convenience messaging. They are best reached via Google Ads and "
     "Instagram, and they convert reliably when the right message is in front of them. "
     "Under a constrained budget, low volume is acceptable when the conversion rate and CPA "
     "are this strong."),

    ("Segment 3 — Young Professionals (Priority 3)",
     "Young Professionals convert at 6% at KES 500 per booking with high volume potential — "
     "the best combination of scale and efficiency outside the top two segments. In Week 3, "
     "C005 (Google Ads, Young Professionals) delivered 120 bookings from 700 clicks at "
     "KES 55,000 — the highest booking volume of any single campaign in the dataset. This "
     "segment is goal-oriented when searching and responds well to clear value propositions "
     "and convenience messaging. They are also the primary retargeting audience on Instagram, "
     "where CTR is 7.5%. The combination of Google Ads for initial capture and Instagram "
     "retargeting for conversion makes this segment viable even under budget pressure."),
]
for title, desc in segments:
    heading(title, level=2)
    body(desc)

# ═════════════════════════════════════════════════════════════════════════════
# 3. TASK 2 — PAID ADVERTISING CHANNEL APPROACH
# ═════════════════════════════════════════════════════════════════════════════
heading("3. Task 2 — Paid Advertising Channel Approach")
body(
    "The strategy uses three channels, each with a clearly defined role. Channel selection "
    "is based on the Week 4 Budget Scenario data, which scores each channel on ROI (1–5) "
    "and risk level, combined with Week 3 campaign performance results."
)

add_table(
    ["Channel", "Role", "ROI Score", "Risk", "Spend Range (KES)", "Primary Segments"],
    [
        ["Google Ads", "Primary conversion",         "5/5", "Low",    "30,000–60,000", "Business Travellers, Families, Young Professionals"],
        ["Instagram",  "Retargeting + mid-funnel",   "4/5", "Medium", "20,000–40,000", "Young Professionals, Leisure Travellers"],
        ["X",          "Reputation management",      "3/5", "Medium", "10,000–20,000", "Business Travellers — trust recovery"],
        ["TikTok",     "Paused",                     "2/5", "High",   "0",             "Excluded under current constraints"],
    ],
    [1.2, 1.8, 0.9, 0.8, 1.5, 2.6]
)

channels = [
    ("Channel 1 — Google Ads: Primary Conversion Channel",
     "Google Ads is the only channel in the dataset that captures users at the moment of "
     "active search intent. A user who searches for a Nairobi–Mombasa bus ticket on Google "
     "has already identified a need and is evaluating options — Travler's ad appears at the "
     "highest point of purchase intent in the entire customer journey. This is why Google "
     "Ads delivers conversion rates of 12–15.7% while social platforms deliver 1.6–6.1%. "
     "Under budget constraints, this is the most defensible channel to Finance because the "
     "link between spend and bookings is direct, measurable, and consistent across two weeks "
     "of data. The recommended spend range is KES 30,000–60,000, targeting Business "
     "Travellers and Young Professionals as primary audiences and Families as secondary. "
     "Search campaigns should be structured around route-specific keywords — 'Nairobi to "
     "Mombasa bus', 'book bus ticket Kenya' — to capture users at the decision stage."),

    ("Channel 2 — Instagram: Retargeting and Mid-Funnel",
     "Instagram serves two functions in this strategy. First, it retargets users who clicked "
     "a Google Ad or visited the Travler site but did not complete a booking — this audience "
     "has already demonstrated intent and requires a lower-cost nudge to convert. Second, it "
     "maintains brand visibility for Young Professionals and Leisure Travellers who are in "
     "the consideration phase but not yet actively searching. Instagram's ROI score of 4/5 "
     "and medium risk rating make it the appropriate secondary channel. In Week 3, Instagram "
     "campaigns delivered conversion rates of 3.3–6.1% — lower than Google Ads but "
     "significantly higher than TikTok. The visual format also supports the trust-rebuilding "
     "messaging that is needed following the Week 3 reputational incident. Retargeting "
     "campaigns should be triggered within 2 hours of a click with no booking, serving a "
     "trust message followed by a price-anchoring offer within 24 hours."),

    ("Channel 3 — X: Reputation Management at Minimum Spend",
     "X is retained in the strategy at minimum spend (KES 10,000) not as a conversion "
     "channel but as a reputation management tool. The Week 3 sentiment analysis showed that "
     "3 of 4 X comments were negative, concentrated on trust and payment themes directly "
     "linked to the viral complaint. Cutting X entirely while this negative sentiment remains "
     "active would leave the reputational damage unaddressed in the platform where it "
     "originated. The minimum spend maintains Travler's presence and supports the pinned "
     "crisis response post recommended in Week 3. Every negative comment on X should receive "
     "a public reply within 2 hours. This is not a growth investment — it is a risk "
     "mitigation cost that protects the return on the Google Ads and Instagram spend by "
     "ensuring that users who search Travler on X after seeing a paid ad do not encounter "
     "unanswered complaints as their first impression."),
]
for title, desc in channels:
    heading(title, level=2)
    body(desc)

# ═════════════════════════════════════════════════════════════════════════════
# 4. TASK 3 — MESSAGING ANGLES
# ═════════════════════════════════════════════════════════════════════════════
heading("4. Task 3 — Messaging Angles by Audience Segment")
body(
    "Three messaging angles have been developed, each tailored to a priority segment's "
    "specific booking behaviour, conversion barrier, and platform context. Each angle "
    "addresses a distinct reason why users in that segment do or do not complete a booking."
)

add_table(
    ["Angle", "Segment", "Core Message", "Barrier Addressed", "Channel"],
    [
        ["Angle 1 — Efficiency & Certainty", "Business Travellers", "Book in 60 seconds. Seat confirmed.",         "Time sensitivity — needs speed and reliability",     "Google Ads, X"],
        ["Angle 2 — Safety & Convenience",   "Families",            "Plan the whole trip. No queues, no stress.",  "Trust — needs safety signals before committing",     "Google Ads, Instagram"],
        ["Angle 3 — Value & Reassurance",    "Young Professionals", "Best price or we match it.",                  "Price perception — hesitates at checkout stage",     "Instagram retargeting"],
    ],
    [1.6, 1.5, 2.0, 2.0, 1.2]
)

angles = [
    ("Angle 1 — Efficiency and Certainty  |  Business Travellers",
     "Business Travellers convert at 14% because they are in an active decision-making state "
     "when they encounter Travler ads. The messaging for this segment does not need to "
     "persuade — it needs to confirm. The core message is speed and reliability: 'Book in "
     "60 seconds. Seat confirmed.' This directly addresses the segment's primary need, which "
     "is certainty. Business Travellers are time-sensitive and brand-agnostic — they will "
     "book the first reliable option that appears. Messaging should lead with the booking "
     "process being fast and the seat being guaranteed, not with price or destination "
     "inspiration. On Google Ads, this translates to ad copy that emphasises instant "
     "confirmation and route availability. On X, it supports the trust recovery narrative "
     "by demonstrating that the platform is reliable and responsive."),

    ("Angle 2 — Safety and Convenience  |  Families",
     "Families convert at 10.5% but require trust signals before committing. The Week 1 "
     "Family Planner segment searched once and booked immediately — but only when the right "
     "conditions were present. The messaging for this segment must answer the unspoken "
     "question: 'Is this safe and easy for my family?' The core message is: 'Plan the whole "
     "trip in one place. No queues, no stress.' Supporting messages should include seat "
     "selection for groups, luggage allowance clarity, punctuality data, and customer "
     "testimonials from other families. On Google Ads, campaigns should target route-specific "
     "searches with family-oriented ad copy. On Instagram, carousel posts featuring real "
     "family testimonials and comfort reassurances convert this segment reliably, as "
     "demonstrated by the Week 2 content performance data."),

    ("Angle 3 — Value and Reassurance  |  Young Professionals",
     "Young Professionals have a 7.5% CTR on Instagram but convert at only 6% — the gap "
     "between click and booking indicates a barrier at the decision stage, not at awareness. "
     "Week 3 sentiment data identified price perception as a key concern: users comparing "
     "Travler's prices to walk-up bus station rates without understanding the value "
     "difference. The core message is: 'Best price or we match it.' This directly addresses "
     "the price hesitation with a guarantee that removes the need to comparison-shop. "
     "Supporting messages should anchor the value proposition: convenience, seat guarantee, "
     "no queuing, and time saved. This angle is deployed primarily through Instagram "
     "retargeting — triggered within 2 hours of a click with no booking — and should be "
     "followed by a limited-time discount offer within 24 hours if the user still has not "
     "converted. The sequence directly mirrors the retargeting recommendation from Week 3."),
]
for title, desc in angles:
    heading(title, level=2)
    body(desc)

# ═════════════════════════════════════════════════════════════════════════════
# 5. TASK 4 — BUDGET ALLOCATION
# ═════════════════════════════════════════════════════════════════════════════
heading("5. Task 4 — Budget Allocation Recommendation")
body(
    "The following allocation is based on a constrained total budget of KES 80,000, "
    "consistent with the minimum-to-mid spend ranges in the Week 4 Budget Scenario data. "
    "The split prioritises channels by ROI score and risk level, with the largest share "
    "going to the channel with the most direct and measurable link to bookings."
)

add_table(
    ["Channel", "Allocation", "KES Amount", "ROI Score", "Risk", "Justification"],
    [
        ["Google Ads", "55%", "KES 44,000", "5/5", "Low",    "Highest conversion rates, lowest CPA, most defensible to Finance"],
        ["Instagram",  "30%", "KES 24,000", "4/5", "Medium", "Retargeting high-CTR non-converters, trust rebuilding post Week 3"],
        ["X",          "15%", "KES 12,000", "3/5", "Medium", "Minimum spend for reputation management — risk mitigation, not growth"],
        ["TikTok",     "0%",  "KES 0",      "2/5", "High",   "Paused — CPA KES 1,000–1,125 indefensible under current constraints"],
    ],
    [1.1, 0.9, 1.1, 0.9, 0.8, 2.6]
)

body(
    "Justification for Kevin R. (Finance): 85% of the total budget is allocated to channels "
    "with ROI scores of 4–5/5 and low-to-medium risk. Google Ads alone accounts for 55% of "
    "spend and has a demonstrated cost per booking of KES 350–455 — meaning KES 44,000 is "
    "projected to generate approximately 97–126 bookings based on Week 3 performance data. "
    "Instagram's KES 24,000 targets users who have already clicked a Travler ad, meaning "
    "the cost of acquisition is lower than a cold audience campaign. The KES 12,000 on X "
    "is a reputation management cost: leaving negative sentiment unaddressed on the platform "
    "where it originated would suppress conversion rates on the other two channels by "
    "undermining trust at the point of decision."
)

body(
    "Justification for Amina K. (Marketing): Travler maintains paid visibility across three "
    "channels. Google Ads keeps the brand present at the moment of highest purchase intent "
    "for the top three segments. Instagram keeps Travler visible to the high-volume Young "
    "Professional segment during the consideration phase and retargets users who showed "
    "intent but did not convert. X keeps the brand's response visible to anyone who "
    "searches Travler following the Week 3 incident. TikTok is paused, not abandoned — "
    "it should be reintroduced once budget pressure eases, repositioned as an awareness "
    "channel with success measured by retargeting pool growth rather than direct bookings."
)

# ═════════════════════════════════════════════════════════════════════════════
# 6. AUDIENCE PERFORMANCE SUMMARY TABLE
# ═════════════════════════════════════════════════════════════════════════════
heading("6. Audience Performance Summary")
body(
    "The table below summarises the full Week 4 audience dataset for reference, showing "
    "all five segments ranked by cost per booking."
)

add_table(
    ["Segment", "Platform Preference", "Avg CTR (%)", "Conv Rate (%)", "Cost per Booking (KES)", "Volume Potential", "Strategy Decision"],
    [
        ["Business Travellers", "Google Ads", "6.8",  "14.0", "350", "Medium", "Priority 1 — full budget allocation"],
        ["Families",            "Google Ads", "5.5",  "10.5", "400", "Low",    "Priority 2 — Google Ads + Instagram"],
        ["Young Professionals", "Instagram",  "7.5",  "6.0",  "500", "High",   "Priority 3 — Google Ads + retargeting"],
        ["Leisure Travellers",  "Instagram",  "7.0",  "4.5",  "600", "Medium", "Secondary — Instagram only, low spend"],
        ["Students",            "TikTok",     "8.2",  "1.5",  "900", "High",   "Excluded — reintroduce when budget allows"],
    ],
    [1.6, 1.4, 0.9, 1.0, 1.5, 1.1, 1.9]
)

# ═════════════════════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ═════════════════════════════════════════════════════════════════════════════
heading("7. Conclusion")
body(
    "The core tension in Week 4 — between Finance's demand for efficiency and Marketing's "
    "need for visibility — is resolved by the data, not by compromise. The data shows that "
    "85% of Travler's marketing budget can be concentrated on two channels (Google Ads and "
    "Instagram) that together cover the full customer journey from search intent to "
    "retargeted conversion, while the remaining 15% on X addresses the reputational risk "
    "that would otherwise suppress the return on the primary spend. This is not a reduced "
    "strategy — it is a more focused one. By removing TikTok from the active spend and "
    "redirecting that budget to Google Ads, the strategy increases projected bookings while "
    "reducing total spend. By maintaining Instagram as a retargeting channel rather than a "
    "cold-audience channel, it lowers the effective cost per acquisition for Young "
    "Professionals. And by keeping X at minimum spend for reputation management, it protects "
    "the trust signals that Business Travellers and Families need before they will commit "
    "to a booking. The three priority segments — Business Travellers, Families, and Young "
    "Professionals — account for the majority of Travler's conversion value in the dataset. "
    "Concentrating spend on them, through the right channels and with the right messages, "
    "gives Travler the clearest path from the current position of budget pressure and "
    "reputational recovery to a position of measurable, defensible growth."
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
set_font(
    p.add_run(
        "Travler  |  Week 4 Paid Advertising Strategy  |  "
        "Dataset: 5 audience segments, 4 channel scenarios  |  "
        "Analysis: audience prioritisation, channel strategy, messaging & budget allocation"
    ),
    8, italic=True
)

doc.save("Travler_Week4_Paid_Advertising_Submission.docx")
print("Saved -> Travler_Week4_Paid_Advertising_Submission.docx")
