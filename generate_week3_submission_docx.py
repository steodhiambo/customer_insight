from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, size, bold=False, color=BLACK, italic=False):
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = color

def heading(text, level=1, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 12, bold=True, color=color)
    return p

def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, 10)
    return p

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.width = Inches(col_widths[i])
        shade_cell(cell, "FFFFFF")
        run = cell.paragraphs[0].add_run(h)
        set_font(run, 9, bold=True, color=BLACK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows):
        bg = "F4F6F8" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.width = Inches(col_widths[c_idx])
            shade_cell(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, 9, color=BLACK)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ═════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(p.add_run("Travler — Week 3 Campaign Performance & Sentiment Analysis"), 22, bold=True)

p2 = doc.add_paragraph()
set_font(p2.add_run(
    "Marketing Team Submission  |  Campaign Review, Customer Sentiment & Brand Perception Strategy"),
    11, italic=True)
p2.paragraph_format.space_after = Pt(10)

# ═════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
heading("1. Executive Summary")
body(
    "Travler is operating in a difficult environment this week. A customer complaint regarding "
    "a double-charge and delayed support response went viral, placing the brand under public "
    "scrutiny at the same time as a multi-channel paid campaign was running. The campaign data "
    "shows a clear performance split: Google Ads is delivering strong conversion rates of "
    "12–15.7% for Business Travellers and Young Professionals, while TikTok is generating "
    "high impressions with conversion rates as low as 1.6%. Instagram sits in the middle, "
    "performing adequately on reach but not efficiently on cost per acquisition. "
    "Simultaneously, sentiment analysis across 10 customer comments reveals that 50% are "
    "negative, concentrated on X (Twitter), and centred on three themes: trust and platform "
    "legitimacy, payment and support failures, and price perception. This report addresses "
    "all four required tasks: four campaign performance insights, identification of high and "
    "underperforming channels, three sentiment themes, and three actionable recommendations."
)

heading("Internal Context", level=2)
body(
    "Amina K. has flagged that the delayed support response has affected public perception and "
    "that the team needs a clear plan to rebuild trust. Diana K. has reported an increase in "
    "inbound complaints and inquiries, consistent with the sentiment data showing unresolved "
    "concerns about payments and refunds. The campaign performance concerns raised by the team "
    "are addressed directly in this report — the data shows the issue is not campaign volume "
    "but channel misalignment, with spend concentrated on platforms that are not converting "
    "at the same rate as Google Ads."
)

add_table(
    ["Metric", "Value", "Implication"],
    [
        ["Total Campaigns Analysed",        "8 (C001–C008)",          "Across Instagram, TikTok, X, and Google Ads"],
        ["Highest Conversion Rate",         "C006 — 15.7%",           "Google Ads, Business Travellers — highest ROI campaign"],
        ["Lowest Conversion Rate",          "C003 / C008 — 1.6%",     "TikTok (both audiences) — reach without intent"],
        ["Best Cost per Acquisition",       "C006 — KES 455",         "Google Ads, Business Travellers — most efficient spend"],
        ["Worst Cost per Acquisition",      "C005 — KES 458",         "Google Ads, Young Professionals — still competitive"],
        ["Overall Sentiment",               "50% Negative",           "Trust and payment concerns dominate — crisis response needed"],
        ["Most Negative Platform",          "X (Twitter)",            "3 of 4 X comments are negative — priority for response"],
        ["Top Negative Theme",              "Trust / Legitimacy",     "Users questioning platform reliability post-viral complaint"],
    ],
    [2.4, 1.8, 2.6]
)

# ═════════════════════════════════════════════════════════════════════════════
# 2. TASK 1 — FOUR KEY CAMPAIGN PERFORMANCE INSIGHTS
# ═════════════════════════════════════════════════════════════════════════════
heading("2. Task 1 — Four Key Campaign Performance Insights")
body(
    "The following four insights summarise the most significant patterns in the Week 3 "
    "campaign dataset, comparing results across all four channels."
)

insights = [
    ("Insight 1 — Google Ads Converts at 8–10× the Rate of TikTok",
     "Google Ads campaigns (C005 and C006) achieved conversion rates of 12.0% and 15.7% "
     "respectively. TikTok campaigns (C003 and C008) achieved 1.6% each. This is not a "
     "marginal difference — it is a structural one. Google Ads reaches users who are "
     "actively searching for travel options, meaning intent is already present at the point "
     "of contact. TikTok reaches users in a passive, entertainment-first mindset. The "
     "platform is not the problem; the mismatch between platform behaviour and conversion "
     "expectation is. TikTok should not be measured against Google Ads on conversion — it "
     "serves a different stage of the customer journey."),

    ("Insight 2 — CTR Does Not Predict Conversion",
     "TikTok's click-through rates (8.0% and 8.3%) are the highest in the dataset, yet its "
     "conversion rates are the lowest. X achieved a 5.0% CTR for C004 but a strong 11.6% "
     "conversion rate. This confirms that clicks measure curiosity, not intent. A user who "
     "taps a TikTok ad is responding to entertainment value; a user who clicks a Google "
     "search ad is responding to a specific need. Campaign optimisation should prioritise "
     "conversion rate and cost per acquisition over CTR as the primary success metric."),

    ("Insight 3 — Business Travellers Deliver the Highest Return on Spend",
     "C006 (Google Ads, Business Travellers) generated 110 bookings from 10,000 impressions "
     "at a cost of KES 50,000 — a cost per acquisition of KES 455. C004 (X, Business "
     "Travellers) generated 70 bookings from 12,000 impressions at KES 30,000 — a CPA of "
     "KES 429. Business Travellers convert at the highest rate across every channel they "
     "appear in. This segment is the most valuable in the dataset and should receive "
     "protected budget allocation regardless of channel decisions made for other segments."),

    ("Insight 4 — Instagram Performs Moderately but Inconsistently Across Audiences",
     "Instagram campaigns ranged from 3.3% conversion (C001, Students) to 6.1% (C002, Young "
     "Professionals) to 4.3% (C007, Leisure Travellers). The platform is mid-tier on both "
     "conversion and cost efficiency. It outperforms TikTok on conversion but significantly "
     "underperforms Google Ads. Instagram's value lies in its visual format and retargeting "
     "capability — it is best used as a mid-funnel channel that bridges TikTok awareness "
     "content and Google Ads conversion, rather than as a primary booking driver."),
]
for title, desc in insights:
    heading(title, level=2)
    body(desc)

# ═════════════════════════════════════════════════════════════════════════════
# 3. TASK 2 — HIGH AND UNDERPERFORMING CHANNELS
# ═════════════════════════════════════════════════════════════════════════════
heading("3. Task 2 — High-Performing and Underperforming Channels")
body(
    "Based on conversion rate, cost per acquisition, and booking volume, two channels are "
    "identified as high-performing and two as underperforming."
)

add_table(
    ["Classification", "Channel / Segment", "Conv Rate", "CPA (KES)", "Key Reason"],
    [
        ["High-Performing", "Google Ads — Business Travellers (C006)", "15.7%", "455",
         "Search intent + high-value audience = highest ROI in dataset"],
        ["High-Performing", "Google Ads — Young Professionals (C005)", "12.0%", "458",
         "Motivated audience, search-based targeting, strong click-to-book ratio"],
        ["Underperforming",  "TikTok — Students (C003)",               "1.6%",  "1,000",
         "Entertainment mindset, price-sensitive, no booking intent at point of contact"],
        ["Underperforming",  "TikTok — Leisure Travellers (C008)",     "1.6%",  "1,125",
         "Passive scrollers, highest impressions but lowest conversion — worst CPA in dataset"],
    ],
    [1.4, 2.2, 0.9, 1.0, 2.3]
)

heading("High-Performing: Google Ads — Business Travellers and Young Professionals", level=2)
body(
    "Both Google Ads campaigns outperform every other channel in the dataset on conversion "
    "rate and cost efficiency. The reason is audience intent: users who search for travel "
    "options on Google are in an active decision-making state. They have already identified "
    "a need and are evaluating options. Travler's ad appears at the moment of highest "
    "purchase intent, which is why the click-to-booking conversion is so much stronger than "
    "on social platforms. Business Travellers in particular are time-sensitive and "
    "brand-agnostic — they will book the first reliable option that appears. Google Ads "
    "captures that moment directly. Young Professionals follow a similar pattern: they are "
    "goal-oriented when searching and respond well to clear value propositions."
)

heading("Underperforming: TikTok — Students and Leisure Travellers", level=2)
body(
    "TikTok's underperformance is not a failure of the platform — it is a failure of "
    "expectation. TikTok users are in a passive, entertainment-first state when they "
    "encounter Travler content. Students are price-sensitive and unlikely to commit to a "
    "purchase from a social feed without a specific incentive. Leisure Travellers are "
    "browsing aspirationally, not planning actively. The result is high impressions, "
    "reasonable CTR, and almost no bookings. The cost per acquisition for TikTok campaigns "
    "is KES 1,000–1,125 — more than double the Google Ads CPA. This does not mean TikTok "
    "should be abandoned; it means TikTok spend should be treated as awareness investment, "
    "not conversion spend, and success should be measured by retargeting pool growth rather "
    "than direct bookings."
)

# ═════════════════════════════════════════════════════════════════════════════
# 4. TASK 3 — THREE KEY SENTIMENT THEMES
# ═════════════════════════════════════════════════════════════════════════════
heading("4. Task 3 — Three Key Sentiment Themes Affecting Trust and Perception")
body(
    "Analysis of 10 customer comments across X, Instagram, and TikTok reveals three "
    "dominant themes. Five comments (50%) are negative, three (30%) are neutral, and two "
    "(20%) are positive. The negative and neutral comments cluster around three concerns "
    "that directly affect willingness to book."
)

add_table(
    ["Theme", "Sentiment", "Platforms", "Comments", "Business Impact"],
    [
        ["Trust & Platform Legitimacy", "Negative / Neutral", "X, Instagram",
         "S002, S007, S010",
         "Users questioning whether Travler is reliable — directly linked to viral complaint"],
        ["Payment & Support Failures",  "Negative",           "X, TikTok",
         "S001, S005, S009",
         "Double-charge, slow support, app crash — operational failures amplifying reputational damage"],
        ["Price Perception",            "Neutral / Negative", "Instagram",
         "S006",
         "Users comparing to walk-up bus station prices — value proposition not communicated"],
    ],
    [1.8, 1.2, 1.2, 1.0, 2.6]
)

heading("Theme 1 — Trust and Platform Legitimacy", level=2)
body(
    "Three comments directly question whether Travler can be trusted as a platform. S002 "
    "states 'Not sure if this platform is reliable.' S007 says 'Will stick to booking "
    "directly.' S010 asks 'Is this platform legit?' These comments are not isolated — they "
    "reflect a broader erosion of confidence triggered by the viral complaint. Users who "
    "were already uncertain about booking online have been given a public reason to remain "
    "hesitant. The trust gap is most acute on X, where the original complaint circulated, "
    "and on Instagram, where Travler's paid campaigns are running. The risk is that "
    "prospective customers who see these comments alongside a paid ad will choose not to "
    "click. Trust must be rebuilt publicly and proactively, not just resolved privately."
)

heading("Theme 2 — Payment and Support Failures", level=2)
body(
    "Three comments describe direct operational failures: S001 reports being charged twice "
    "with no quick response; S005 states that support takes too long; S009 reports the app "
    "freezing during booking. These are not perception problems — they are product and "
    "process failures that are being aired publicly. Each unresolved complaint on a public "
    "platform functions as a negative advertisement. The support response time issue is "
    "particularly damaging because it compounds the original payment error: the customer "
    "was not only charged incorrectly but was left without resolution. Diana K.'s report of "
    "increased inbound complaints is consistent with this pattern — users who do not receive "
    "a public response escalate to direct contact. Faster public responses would reduce "
    "both the reputational damage and the inbound support volume."
)

heading("Theme 3 — Price Perception", level=2)
body(
    "S006 states that 'prices seem higher than bus station.' This is a single comment but "
    "it represents a structural challenge: Travler is competing against the mental benchmark "
    "of walk-up bus station pricing, which users perceive as the default and cheapest "
    "option. Without a visible value proposition — convenience, safety, seat guarantee, "
    "time saving — the platform appears to charge a premium for no clear reason. This "
    "perception is particularly damaging for the Student and Leisure segments, who are "
    "already price-sensitive. A Best Price Guarantee message or a clear 'why book online' "
    "explanation on the booking page and in campaign creative would directly address this "
    "concern and reduce the hesitation that drives repeat searching without conversion."
)

# ═════════════════════════════════════════════════════════════════════════════
# 5. TASK 4 — THREE RECOMMENDATIONS
# ═════════════════════════════════════════════════════════════════════════════
heading("5. Task 4 — Three Recommendations")
body(
    "The following three recommendations address both campaign performance and brand "
    "perception. Each recommendation is grounded in the data and designed to produce a "
    "measurable outcome."
)

recs = [
    ("Recommendation 1 — Reallocate 20–30% of TikTok Budget to Google Ads",
     "The data makes a clear case for budget reallocation. TikTok's CPA is KES 1,000–1,125. "
     "Google Ads CPA is KES 455–458. Every KES 1,000 spent on TikTok generates approximately "
     "one booking. The same spend on Google Ads generates approximately two. Shifting 20–30% "
     "of TikTok spend to Google Ads — specifically targeting Business Travellers and Young "
     "Professionals — would increase total bookings without increasing total spend. TikTok "
     "should remain in the media mix but be repositioned as a brand awareness channel. "
     "Success on TikTok should be measured by retargeting audience growth, not direct "
     "bookings. Users who engage with TikTok content should be served Google Ads or "
     "Instagram retargeting ads within 24–48 hours to complete the conversion journey."),

    ("Recommendation 2 — Publish a Public Crisis Response on X Within 24 Hours",
     "The trust and payment themes in the sentiment data are concentrated on X, which is "
     "also the platform where the original viral complaint circulated. A public response is "
     "required — not a private resolution. The response should acknowledge the payment "
     "error directly, confirm it has been resolved, state the steps taken to prevent "
     "recurrence, and commit to a specific support response time (recommended: 2 hours "
     "during business hours). This post should be pinned to the Travler X profile. "
     "Every negative comment currently visible on X should receive a personalised public "
     "reply within 24 hours of this submission. Users who see a brand respond quickly and "
     "transparently to a complaint are significantly more likely to trust the platform than "
     "users who see silence. The goal is not to delete the negative narrative but to "
     "replace it with a visible demonstration of accountability."),

    ("Recommendation 3 — Deploy a Retargeting Sequence for High-CTR Non-Converters",
     "Students and Young Professionals are clicking on Travler ads — CTR for these segments "
     "ranges from 7.2% to 8.0% — but not completing bookings. The gap between click and "
     "conversion indicates that the barrier is not awareness or interest but something at "
     "the decision stage: price uncertainty, trust hesitation, or checkout friction. A "
     "retargeting sequence should be triggered for any user who clicks a Travler ad but "
     "does not complete a booking within 2 hours. The sequence should include three "
     "touchpoints: first, a trust message addressing the platform's reliability and payment "
     "security; second, a price-anchoring message that explains the value of booking online "
     "versus walk-up (convenience, seat guarantee, no queuing); third, a limited-time "
     "discount offer of 5–10% to create urgency. This sequence directly addresses all three "
     "sentiment themes — trust, payment confidence, and price perception — and targets the "
     "segment with the highest volume of unconverted clicks."),
]
for title, desc in recs:
    heading(title, level=2)
    body(desc)

# ═════════════════════════════════════════════════════════════════════════════
# 6. SENTIMENT SUMMARY TABLE
# ═════════════════════════════════════════════════════════════════════════════
heading("6. Customer Sentiment Summary")
body(
    "The table below lists all ten sentiment comments with their platform, tone, and theme "
    "for reference."
)

add_table(
    ["Comment ID", "Platform", "Sentiment", "Theme", "Summary"],
    [
        ["S001", "X",         "Negative", "Payment Issue",    "Charged twice and no quick response"],
        ["S002", "Instagram", "Negative", "Trust",            "Not sure if this platform is reliable"],
        ["S003", "X",         "Neutral",  "Inquiry",          "How do refunds work?"],
        ["S004", "Instagram", "Positive", "Convenience",      "Booking process is easy"],
        ["S005", "TikTok",    "Negative", "Customer Service", "Support takes too long"],
        ["S006", "Instagram", "Neutral",  "Price",            "Prices seem higher than bus station"],
        ["S007", "X",         "Negative", "Trust",            "Will stick to booking directly"],
        ["S008", "Instagram", "Positive", "Discounts",        "Got a good deal on tickets"],
        ["S009", "TikTok",    "Negative", "Experience",       "App froze during booking"],
        ["S010", "X",         "Neutral",  "Clarification",    "Is this platform legit?"],
    ],
    [1.0, 1.0, 1.0, 1.4, 3.4]
)

# ═════════════════════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ═════════════════════════════════════════════════════════════════════════════
heading("7. Conclusion")
body(
    "Travler faces two simultaneous challenges this week: a campaign performance gap driven "
    "by channel misalignment, and a reputational challenge driven by a viral complaint and "
    "its aftermath. The data shows these challenges are connected. Users who encounter "
    "negative sentiment about Travler on X before or after seeing a paid ad are less likely "
    "to convert — the trust deficit created by the complaint is actively suppressing the "
    "return on campaign spend. Addressing both issues together is therefore more valuable "
    "than treating them separately. The three recommendations in this report — reallocating "
    "budget to Google Ads, publishing a public crisis response on X, and deploying a "
    "retargeting sequence for non-converters — are designed to work as a coordinated "
    "response. Budget reallocation improves conversion efficiency. The crisis response "
    "rebuilds the trust that makes conversion possible. The retargeting sequence captures "
    "the users who are interested but hesitant, addressing the price and trust barriers "
    "that the sentiment data has identified. Together, these actions give Travler a "
    "structured path from the current situation — high reach, low conversion, damaged "
    "trust — to a position of improved performance and rebuilt credibility."
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
set_font(
    p.add_run(
        "Travler  |  Week 3 Campaign & Sentiment Analysis  |  "
        "Dataset: 8 campaigns, 10 sentiment comments, 4 channels  |  "
        "Analysis: campaign performance, customer sentiment & brand perception"
    ),
    8, italic=True
)

doc.save("Travler_Week3_Campaign_Sentiment_Submission.docx")
print("Saved → Travler_Week3_Campaign_Sentiment_Submission.docx")
