from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, size, bold=False, color=BLACK, italic=False):
    run.font.size      = Pt(size)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = color

def heading(text, level=1, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 16 if level == 1 else 12, bold=True, color=color)
    return p

def body(text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_font(run, 10)
    return p

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_table(headers, rows, col_widths):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.width = Inches(col_widths[i])
        shade_cell(cell, "FFFFFF")
        run = cell.paragraphs[0].add_run(h)
        set_font(run, 9, bold=True, color=BLACK)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows):
        bg = "F4F6F8" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.width = Inches(col_widths[c_idx])
            shade_cell(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, 9, color=BLACK)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ═════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(p.add_run("Travler — Week 2 Content Strategy & Performance Analysis"), 22, bold=True)

p2 = doc.add_paragraph()
set_font(p2.add_run("Marketing Team Submission  |  Content Planning, Audience Insights & Campaign Strategy"),
         11, italic=True)
p2.paragraph_format.space_after = Pt(10)

# ═════════════════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
heading("1. Executive Summary")
body(
    "Travler's content activity is generating reach but not reliably converting that reach into "
    "bookings. An analysis of 15 posts across Instagram, TikTok, and X reveals a clear pattern: "
    "promotional and booking-focused content on Instagram — particularly Reels — drives the "
    "highest number of completed bookings, while TikTok builds awareness without converting, "
    "and X underperforms across every metric. The data also confirms a persistent gap between "
    "the platform's most active audience segments (Students and Young Professionals) and actual "
    "booking behaviour, a pattern first identified in Week 1. This report translates those "
    "findings into a structured content strategy, a two-week content calendar, and sample posts "
    "that align tone, platform, and intent."
)

heading("Internal Context", level=2)
body(
    "This strategy has been developed with three internal concerns in mind. Amina K. raised "
    "that content is reactive rather than strategic — the three content themes and two-week "
    "calendar in this submission directly address that by giving the team a planned, "
    "intentional structure to work from. Kevin R. questioned whether content is delivering "
    "measurable value — every post in the calendar is mapped to a business outcome, and the "
    "Book It theme is specifically designed to generate bookings that can be tracked against "
    "content spend. Diana K. noted that customer questions on social platforms go unanswered "
    "— the Know Before You Go theme converts those questions into proactive content, reducing "
    "the support burden while building trust with users who are close to booking."
)

add_table(
    ["Metric", "Value", "Implication"],
    [
        ["Total Posts Analysed",              "15",                    "Across Instagram, TikTok, and X"],
        ["Highest Booking Post",              "P009 — 60 bookings",    "Instagram Reel, Discount Offer theme"],
        ["Lowest Booking Post",               "P014 — 8 bookings",     "X Text, Customer Support Info"],
        ["Top Engagement Format",             "Video / Reel",          "Highest avg engagement rate across all formats"],
        ["Platform with Most Bookings",       "Instagram",             "Reels and Carousels drive direct conversion"],
        ["Platform with Lowest Booking Rate", "X",                     "Minimal reach, engagement, and conversion"],
        ["TikTok Engagement vs Bookings",     "High reach, low conv.", "Awareness only — not a booking channel"],
    ],
    [2.4, 1.8, 2.6]
)

# ═════════════════════════════════════════════════════════════════════════════
# 2. KEY PATTERNS FROM CONTENT PERFORMANCE DATA
# ═════════════════════════════════════════════════════════════════════════════
heading("2. Key Patterns from Content Performance Data")
body(
    "Five patterns emerge from the content performance dataset that explain what is working, "
    "what is not, and where the most significant gaps exist between content activity and "
    "business outcomes."
)

patterns = [
    ("Pattern 1 — Reels and Videos Dominate Reach",
     "TikTok posts reached 18,000–21,000 impressions per post. Instagram Reels reached "
     "14,000–16,000. X text posts reached 5,000–8,000. Video is the highest-reach format "
     "by a significant margin, and any content strategy that does not prioritise video will "
     "underperform on reach."),
    ("Pattern 2 — Promotional Content Converts; Entertainment Content Does Not",
     "P009 (Instagram Reel, Discount Offer) generated 60 bookings — the highest in the dataset. "
     "P015 (Instagram Reel, Booking Guide) generated 50. By contrast, P012 (TikTok, "
     "Entertainment) generated 2,600 likes and 170 comments but only 22 bookings. High "
     "engagement does not equal high conversion. Content with a direct call to action and a "
     "clear value proposition converts; entertainment content builds awareness only."),
    ("Pattern 3 — TikTok Is an Awareness Channel, Not a Booking Channel",
     "TikTok consistently delivers the highest impressions and engagement rates but the lowest "
     "booking numbers relative to reach. This is consistent with platform behaviour — TikTok "
     "users are in a passive consumption mindset, not a decision-making one. Using TikTok to "
     "drive direct bookings is a misalignment of platform and intent."),
    ("Pattern 4 — X Is Not Delivering Value",
     "Across all five metrics — impressions, likes, comments, shares, and bookings — X posts "
     "rank last. The highest-performing X post (P003, Route Info) generated 8,000 impressions "
     "and 20 bookings. The lowest (P014, Customer Support Info) generated 5,000 impressions "
     "and 8 bookings. The platform is not reaching Travler's audience at meaningful scale and "
     "should be deprioritised for content creation."),
    ("Pattern 5 — Trust-Based Content Converts Quietly but Consistently",
     "P007 (Customer Testimonial, Instagram Image) generated 35 bookings from 9,000 impressions "
     "— a booking rate of 0.39%, which is competitive with the top Reel posts. P011 (Route "
     "Highlight) generated 38 bookings from 11,000 impressions. Trust-building content does "
     "not generate viral engagement but it converts reliably, particularly for Family and "
     "Business segments identified in Week 1."),
]
for title, desc in patterns:
    heading(title, level=2)
    body(desc)

# ═════════════════════════════════════════════════════════════════════════════
# 3. CONTENT THEMES
# ═════════════════════════════════════════════════════════════════════════════
heading("3. Content Themes")
body(
    "Three content themes have been defined to structure Travler's output. Each theme serves "
    "a distinct purpose in the customer journey and maps directly to the audience segments and "
    "behavioural patterns identified across Weeks 1 and 2."
)

add_table(
    ["Theme", "Purpose", "Primary Formats", "Target Segments", "Business Outcome"],
    [
        ["Book It",
         "Drive direct bookings through promotional content with a clear CTA",
         "Instagram Reels, Carousels",
         "All segments — especially Leisure and Students",
         "Increased conversion rate"],
        ["Know Before You Go",
         "Build trust and answer customer questions through educational content",
         "Instagram Carousels, Images",
         "Business Travellers, Families",
         "Reduced support queries, higher trust"],
        ["Travler Life",
         "Build brand awareness and top-of-funnel reach through entertaining content",
         "TikTok Videos, Instagram Reels",
         "Students, Young Professionals",
         "Audience growth, retargeting pool"],
    ],
    [1.2, 2.0, 1.6, 1.8, 1.8]
)

themes = [
    ("Theme 1 — Book It",
     "This theme covers all content designed to move a user from consideration to booking. "
     "It includes discount offers, limited-seat alerts, booking guides, and fare promotions. "
     "The data shows this theme generates the highest bookings per post when delivered via "
     "Instagram Reels with a direct link in bio or swipe-up CTA. Urgency language — 'seats "
     "filling fast', 'this weekend only' — is consistent with the behaviour of converters "
     "identified in Week 1, who search once and commit quickly. This theme directly supports "
     "Kevin R.'s requirement for content that demonstrates measurable ROI."),
    ("Theme 2 — Know Before You Go",
     "This theme covers educational and trust-building content: route highlights, travel tips, "
     "customer testimonials, and safety or comfort reassurances. It addresses the gap Diana K. "
     "identified — users asking questions on social platforms that go unanswered. By surfacing "
     "answers proactively as content, Travler reduces friction at the decision stage and builds "
     "the confidence that Family and Business segments need before booking. Testimonials in "
     "particular convert at a rate comparable to promotional Reels, with lower production cost."),
    ("Theme 3 — Travler Life",
     "This theme covers entertainment and lifestyle content: travel hacks, day-in-the-life "
     "videos, and user-generated style content. It is designed for TikTok and targets Students "
     "and Young Professionals — the two segments with the highest search activity and lowest "
     "conversion. The goal is not direct booking but audience growth and retargeting: users "
     "who engage with Travler Life content are added to a retargeting audience and served "
     "Book It content within 24–48 hours. This two-step approach reflects the platform "
     "behaviour of these segments rather than fighting against it."),
]
for title, desc in themes:
    heading(title, level=2)
    body(desc)

# ═════════════════════════════════════════════════════════════════════════════
# 4. TWO-WEEK CONTENT CALENDAR
# ═════════════════════════════════════════════════════════════════════════════
heading("4. Two-Week Content Calendar")
body(
    "The calendar below covers 10 posts across two weeks. Posts are distributed across "
    "Instagram and TikTok, with X excluded from primary content creation. Each post is "
    "mapped to a theme, segment, and intended business outcome."
)

add_table(
    ["Day", "Platform", "Format", "Theme", "Segment", "Content Idea"],
    [
        ["Mon Wk 1", "Instagram", "Reel",     "Book It",            "All",                 "Flash fare — Nairobi–Mombasa weekend deal, seats limited"],
        ["Tue Wk 1", "TikTok",    "Video",    "Travler Life",       "Students",            "'Pack in 5 minutes' travel hack for the Kisumu route"],
        ["Wed Wk 1", "Instagram", "Carousel", "Know Before You Go", "Young Professionals", "3 things to check before booking a business trip"],
        ["Thu Wk 1", "Instagram", "Image",    "Know Before You Go", "Families",            "Customer testimonial — family trip to Nakuru, safety focus"],
        ["Fri Wk 1", "Instagram", "Reel",     "Book It",            "Young Professionals", "Booking guide — how to confirm your seat in under 2 minutes"],
        ["Mon Wk 2", "TikTok",    "Video",    "Travler Life",       "Students",            "Day-in-the-life on the Nairobi–Kisumu route"],
        ["Tue Wk 2", "Instagram", "Carousel", "Know Before You Go", "Business",            "Route highlight — Nairobi–Eldoret, schedule and comfort"],
        ["Wed Wk 2", "Instagram", "Reel",     "Book It",            "All",                 "Limited seats promo — countdown urgency, direct booking CTA"],
        ["Thu Wk 2", "TikTok",    "Video",    "Travler Life",       "Young Professionals", "'Why I switched to Travler' — UGC-style testimonial"],
        ["Fri Wk 2", "Instagram", "Image",    "Know Before You Go", "Families",            "Safety and comfort reassurance — what families can expect"],
    ],
    [0.9, 1.0, 0.9, 1.5, 1.5, 2.6]
)

body(
    "Note on X: X is excluded from the primary content calendar based on performance data. "
    "If the platform must be maintained, limit activity to customer support replies and "
    "announcement reposts only. Do not invest original content creation resource in X at "
    "this stage."
)

# ═════════════════════════════════════════════════════════════════════════════
# 5. SAMPLE POSTS
# ═════════════════════════════════════════════════════════════════════════════
heading("5. Sample Posts")
body(
    "The two sample posts below demonstrate how Travler's tone and messaging can be applied "
    "across different themes and segments. Both posts include a platform rationale, format "
    "guidance, caption, and intended action."
)

heading("Sample Post 1 — Instagram Reel  |  Book It Theme  |  All Segments", level=2)
body(
    "Platform: Instagram\n"
    "Format: Reel (15–30 seconds)\n"
    "Theme: Book It\n"
    "Target Segment: All — with particular relevance to Leisure and Student segments who "
    "browse without committing\n\n"
    "Caption:\n"
    "\"Nairobi to Mombasa this weekend? 🌊 Seats are filling fast — grab yours before "
    "they're gone. Tap the link in bio to book in under 2 minutes. #Travler "
    "#WeekendGetaway #NairobiToMombasa\"\n\n"
    "Visual Direction: Fast-cut footage of the route, destination arrival, and a clean "
    "booking screen. End frame shows seat availability counter ticking down.\n\n"
    "Intent: Urgency combined with a frictionless CTA. Directly addresses the hesitation "
    "behaviour of non-converters identified in Week 1 — users who search repeatedly without "
    "a reason to stop. The countdown mechanic provides that reason."
)

heading("Sample Post 2 — Instagram Carousel  |  Know Before You Go Theme  |  Families", level=2)
body(
    "Platform: Instagram\n"
    "Format: Carousel (4 slides)\n"
    "Theme: Know Before You Go\n"
    "Target Segment: Families — 100% conversion rate in Week 1 data, motivated by trust "
    "and certainty\n\n"
    "Caption:\n"
    "\"Travelling with the family? Here's what Travler parents actually say 👨‍👩‍👧 Swipe to "
    "see real experiences on the Nairobi–Nakuru route — comfort, safety, and zero stress. "
    "Book your seats today. #FamilyTravel #Travler #TravelKenya\"\n\n"
    "Slide Structure:\n"
    "  Slide 1 — Hook: 'What do families say about Travler?'\n"
    "  Slide 2 — Testimonial quote from a parent, photo of family on route\n"
    "  Slide 3 — Key comfort and safety features (luggage space, seat comfort, punctuality)\n"
    "  Slide 4 — CTA: 'Book the Nakuru route — link in bio'\n\n"
    "Intent: Addresses Diana K.'s observation that customer questions go unanswered on "
    "social platforms. By surfacing real testimonials as structured content, Travler answers "
    "the trust question before it is asked and removes the final barrier for Family segment "
    "users who are already close to booking."
)

# ═════════════════════════════════════════════════════════════════════════════
# 6. AUDIENCE SEGMENTS — CONTENT MAPPING
# ═════════════════════════════════════════════════════════════════════════════
heading("6. Audience Segments — Content Mapping")
body(
    "The four audience segments from the Week 2 dataset map directly onto the three content "
    "themes. The table below shows which content type each segment responds to and what "
    "business outcome is expected."
)

add_table(
    ["Segment", "Age", "Device", "Content Preference", "Recommended Theme", "Expected Outcome"],
    [
        ["Students",            "18–24", "Mobile",  "Entertainment + Discounts",  "Travler Life → Book It retarget", "Convert 0% segment via two-step funnel"],
        ["Young Professionals", "25–34", "Mobile",  "Travel Tips + Convenience",  "Know Before You Go + Book It",    "Increase moderate conversion rate"],
        ["Business Travellers", "25–44", "Desktop", "Efficiency + Reliability",   "Know Before You Go",              "Protect and grow high-value segment"],
        ["Families",            "35+",   "Desktop", "Trust + Safety",             "Know Before You Go + Book It",    "Maintain 100% conversion, increase frequency"],
    ],
    [1.4, 0.6, 0.8, 1.8, 1.8, 2.0]
)

# ═════════════════════════════════════════════════════════════════════════════
# 7. RECOMMENDATIONS
# ═════════════════════════════════════════════════════════════════════════════
heading("7. Recommendations")

recs = [
    ("1. Prioritise Instagram Reels for Conversion Content",
     "The data is unambiguous — Instagram Reels with Discount Offer and Booking Guide themes "
     "generate the most bookings per post. Allocate the majority of content production budget "
     "to this format and theme combination. Aim for a minimum of two Reels per week during "
     "the two-week calendar period."),
    ("2. Use TikTok for Awareness Only — Not Direct Conversion",
     "TikTok reaches the largest audience but converts at the lowest rate. Reframe TikTok's "
     "role as a top-of-funnel awareness channel. Build a retargeting audience from TikTok "
     "engagements and serve those users Book It content on Instagram within 24–48 hours. "
     "Do not measure TikTok success by bookings — measure it by audience growth and "
     "retargeting pool size."),
    ("3. Deprioritise X for Content Creation",
     "X is not delivering meaningful reach, engagement, or bookings for Travler's audience. "
     "Redirect the time and resource currently spent on X content to Instagram Reels and "
     "TikTok. Maintain X only for customer support responses and critical announcements."),
    ("4. Fix the Post-Click Journey",
     "Several posts show high click numbers that do not translate into proportional bookings, "
     "suggesting friction between the click and the completed booking. Audit the mobile "
     "landing page and checkout flow — this is consistent with the mobile conversion gap "
     "identified in Week 1 and remains the single highest-impact technical fix available."),
    ("5. Create Segment-Specific Content Tracks",
     "Students need Discount Reels and Entertainment content that feeds into a retargeting "
     "funnel. Business Travellers need Efficiency Carousels and Route Highlights. Families "
     "need Trust and Testimonial posts. Producing generic content that attempts to speak to "
     "all segments simultaneously will continue to underperform. Segment-specific tracks "
     "require more planning but the Week 1 and Week 2 data both confirm the return is "
     "significantly higher."),
]
for title, desc in recs:
    heading(title, level=2)
    body(desc)

# ═════════════════════════════════════════════════════════════════════════════
# 8. CONCLUSION
# ═════════════════════════════════════════════════════════════════════════════
heading("8. Conclusion")
body(
    "Travler's content challenge is not a volume problem — 15 posts across three platforms "
    "represents a reasonable output. The problem is alignment: content is not consistently "
    "matched to the right platform, format, theme, or audience segment. The Week 2 data "
    "makes the solution clear. Instagram Reels with promotional themes convert. TikTok builds "
    "reach but requires a retargeting bridge to convert. X is not contributing to business "
    "outcomes and should be deprioritised. Trust-based content converts quietly but reliably "
    "for the highest-value segments. The two-week content calendar, three content themes, and "
    "sample posts in this submission provide a structured, evidence-based starting point for "
    "moving Travler's content from reactive to intentional — and from engagement metrics to "
    "booking outcomes."
)

# ── Footer ────────────────────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
set_font(
    p.add_run("Travler  |  Week 2 Content Strategy  |  Dataset: 15 posts, 4 audience segments  |  Analysis: content performance & campaign planning"),
    8, italic=True
)

doc.save("Travler_Week2_Content_Strategy_Submission.docx")
print("Saved → Travler_Week2_Content_Strategy_Submission.docx")
