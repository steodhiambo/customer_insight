import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_submission():
    doc = docx.Document()

    # Title
    title = doc.add_heading('Travler Week 6: Customer Retention Strategy', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Task 1
    doc.add_heading('Task 1', level=1)
    
    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run("High Engagement (Active): ").bold = True
    p1.add_run("Business Travellers and Young Professionals.\n")
    p1.add_run("Behavior: ").italic = True
    p1.add_run("High booking frequency (3–8 trips) with short intervals (5–15 days).\n")
    p1.add_run("Engagement: ").italic = True
    p1.add_run("Highly responsive, showing open rates up to 40% for morning communications.")

    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run("Medium Engagement (At-Risk): ").bold = True
    p2.add_run("Leisure Travellers and Families.\n")
    p2.add_run("Behavior: ").italic = True
    p2.add_run("Inconsistent usage; typically 2–3 bookings, but inactive for the last 20–30 days.\n")
    p2.add_run("Engagement: ").italic = True
    p2.add_run("Moderate responsiveness (20–30% open rates) during afternoon hours.")

    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run("Low Engagement (Inactive): ").bold = True
    p3.add_run("Students.\n")
    p3.add_run("Behavior: ").italic = True
    p3.add_run("Single-use customers who have remained inactive for 45–90 days.\n")
    p3.add_run("Engagement: ").italic = True
    p3.add_run("Minimal interaction (8–12% open rates), peaking slightly in the evening.")

    # Task 2
    doc.add_heading('Task 2', level=1)
    
    p4 = doc.add_paragraph(style='List Bullet')
    p4.add_run("Campaign 1: Semester Break Travel Guide\n").bold = True
    p4.add_run("Target Segment: ").bold = True
    p4.add_run("Inactive (Students)\n")
    p4.add_run("Objective: ").bold = True
    p4.add_run("Drive the critical second booking.\n")
    p4.add_run("Focus: ").bold = True
    p4.add_run("Sent during evening hours, this campaign highlights curated budget-friendly routes and 'student-exclusive off-peak fares' to address their documented price sensitivity without relying on generic mass-discounts.")

    p5 = doc.add_paragraph(style='List Bullet')
    p5.add_run("Campaign 2: The Fast-Track Experience\n").bold = True
    p5.add_run("Target Segment: ").bold = True
    p5.add_run("Active (Business Travellers)\n")
    p5.add_run("Objective: ").bold = True
    p5.add_run("Reinforce loyalty and maximize lifetime value.\n")
    p5.add_run("Focus: ").bold = True
    p5.add_run("Sent in the morning, the message introduces a new one-tap rebooking feature and priority support, directly appealing to their need for speed and convenience.")

    # Task 3
    doc.add_heading('Task 3', level=1)

    doc.add_heading('Email 1: Inactive Segment (Students)', level=2)
    p7 = doc.add_paragraph()
    p7.add_run("Subject line: ").bold = True
    p7.add_run("Unlock your exclusive student fares 🎒")
    doc.add_paragraph("Core message:\nHi [Name],\nIt’s been a while since your last trip! We know student life is busy and budgets matter, so we’ve curated budget-friendly options just for you.\n\nWe’ve unlocked exclusive student off-peak fares on your favourite routes. Whether you're heading home for the weekend or exploring a new city, we’ve got you covered with flexible booking options. Log in to explore your new student portal and start planning!")

    doc.add_heading('Email 2: Active Segment (Business Travellers)', level=2)
    p8 = doc.add_paragraph()
    p8.add_run("Subject line: ").bold = True
    p8.add_run("Streamline your upcoming work trips 🚀")
    doc.add_paragraph("Core message:\nHi [Name],\nWe noticed you’ve been travelling frequently for work lately. Thank you for consistently choosing Travler!\n\nTo make your busy schedule a little easier, we’ve introduced a one-tap rebooking feature for your frequent routes. Plus, as a top-tier traveller, you now have priority access to our dedicated support line.\n\nBook your next work trip in seconds and focus on what really matters. Safe travels!")

    # Task 4
    doc.add_heading('Task 4', level=1)
    
    doc.add_heading('Encourage repeat bookings', level=2)
    doc.add_paragraph("The approach targets specific barriers to re-booking. For Inactive segments (Students), providing exclusive off-peak fares addresses their price sensitivity, lowering the financial hurdle for a second booking. For Active segments (Business Travellers), introducing one-tap rebooking removes friction, transforming frequent travel from a chore into a seamless habit.")
    
    doc.add_heading('Improve customer engagement', level=2)
    doc.add_paragraph("Engagement improves when communication provides relevant value rather than generic noise. By sending At-Risk segments (Families) destination inspiration during afternoon planning phases—before they churn—we naturally capture their interest and keep Travler top-of-mind during their consideration phase.")

    doc.add_heading('Refer to customer behaviour patterns', level=2)
    doc.add_paragraph("These strategies are rooted in data. Students exhibit high price sensitivity and minimal engagement (8-12% open rates), requiring direct financial incentives. Conversely, Business Travellers book every 5-15 days and engage heavily (40% open rates). This proves they value speed and convenience over pricing. Aligning our messaging with these specific behaviors drives long-term retention.")

    # Save
    doc.save('Steph_Travler_Week6.docx')
    print("Document saved successfully as Steph_Travler_Week6.docx")

if __name__ == "__main__":
    create_submission()
