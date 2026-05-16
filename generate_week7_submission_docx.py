import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_submission():
    doc = docx.Document()

    # Title
    title = doc.add_heading('Travler Week 7: Brand Positioning & Partnerships', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ---------------- Task 1 ----------------
    doc.add_heading('Task 1: Positioning Gaps', level=1)
    doc.add_paragraph(
        "Based on the brand perception and competitor data, three gaps stand out in how Travler is currently perceived:"
    )

    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run("Trust deficit vs. traditional channels (gap -2). ").bold = True
    p1.add_run(
        "Travler scores 3/5 on trust against 5/5 for bus operators and station agents. Customers still see direct operators and in-person agents as safer, particularly on familiar routes, which slows adoption despite the platform working well."
    )

    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run("Low brand recognition (gap -2). ").bold = True
    p2.add_run(
        "Travler scores 2/5 on brand recognition vs. 4/5 for top competitors. The brand is functionally capable but not yet top-of-mind, so customers default to operators or aggregators they already know."
    )

    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run("Weak price-value perception (gap -1). ").bold = True
    p3.add_run(
        "Discount-led players such as QuickTicket KE are seen as cheaper, even where Travler is competitive. Travler is not clearly communicating value, leaving room for price-led rivals to define the conversation."
    )

    doc.add_paragraph(
        "Underlying these gaps is a messaging issue: Travler talks about convenience, which competitors now match. Its real advantage — the widest route and operator comparison (4/5 vs. 3/5) — is under-communicated, so the brand is not yet associated with a distinctive, ownable benefit."
    )

    # ---------------- Task 2 ----------------
    doc.add_heading('Task 2: Positioning Statement', level=1)
    p = doc.add_paragraph()
    p.add_run("Positioning direction: ").bold = True
    p.add_run(
        "Travler is the trusted route-comparison platform for everyday travellers — the easiest way to compare verified bus operators, prices and schedules in one place, and book the right trip with confidence."
    )
    doc.add_paragraph(
        "This shifts Travler from a generic “convenient booking app” to owning a specific, defensible space: trusted comparison across operators. It directly leverages the route-variety strength (+1), counters the trust gap by emphasising verified operators, and differentiates from single-operator sites, discount aggregators and offline agents."
    )

    # ---------------- Task 3 ----------------
    doc.add_heading('Task 3: Partnership Opportunities', level=1)

    p4 = doc.add_paragraph(style='List Bullet')
    p4.add_run("1. Bus operator co-marketing (KES 120,000 — High reach, High alignment). ").bold = True
    p4.add_run(
        "Joint campaigns with established bus operators where Travler is promoted on operator channels (ticket counters, buses, social pages) and operators are featured as “verified partners” on Travler. "
    )
    p4_v = doc.add_paragraph()
    p4_v.add_run("Value: ").bold = True
    p4_v.add_run(
        "Directly closes the trust gap (-2) by borrowing operator credibility, reinforces the route-comparison positioning, and exposes Travler to high-intent travellers already at the point of booking. Strong brand alignment because both parties sell the same journey."
    )

    p5 = doc.add_paragraph(style='List Bullet')
    p5.add_run("2. University travel clubs (KES 80,000 — Medium reach, High alignment). ").bold = True
    p5.add_run(
        "Partner with student travel clubs and SRCs across key campuses for branded route guides, term-break booking drives, referral codes and on-campus activations. "
    )
    p5_v = doc.add_paragraph()
    p5_v.add_run("Value: ").bold = True
    p5_v.add_run(
        "Builds brand recognition (gap -2) with a digitally native, frequent-travel segment at low cost. Students are repeat users on intercity routes and become long-term advocates, supporting awareness now and loyalty later."
    )

    # ---------------- Task 4 ----------------
    doc.add_heading('Task 4: Partnership Prioritisation', level=1)
    p6 = doc.add_paragraph()
    p6.add_run("Prioritise: Bus operator co-marketing (KES 120,000). ").bold = True
    p6.add_run(
        "It addresses the trust gap (-2) at the broadest reach — all customer segments — rather than spending KES 80,000 on a university programme limited to students (0–1.5% historical conversion). Associating Travler with operators customers already trust reinforces the route-comparison positioning. Phase 1 (Months 1–3): launch bus operator co-marketing on 2–3 key routes. Phase 2 (Months 4–6): add university programme for awareness-building, using credibility from operator deals to strengthen activation."
    )

    # Save
    doc.save('Steph_Travler_Week7.docx')
    print("Document saved successfully as Steph_Travler_Week7.docx")


if __name__ == "__main__":
    create_submission()
